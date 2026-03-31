"""Tests for generate_factsheet() comparing output against expected DOCX files.

Compares all table cell text and inline image presence between generated
and expected factsheets for both EN and CN languages.
"""

import io
import hashlib
import re
from datetime import datetime
from pathlib import Path

import pytest
from docx import Document

from fofproject.fund import input_monthly_returns
from fofproject.document import generate_factsheet
from fofproject.paths import TEMPLATE_DIR

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
FIXTURES_DIR = PROJECT_ROOT / "tests" / "fixtures" / "documents"
EXPECTED_DIR = FIXTURES_DIR / "expected_outcome"
INPUT_DIR = PROJECT_ROOT / "input"


@pytest.fixture(scope="module")
def funds():
    """Load all funds + benchmarks once for the module."""
    return input_monthly_returns(
        str(INPUT_DIR / "RETURN DATA.csv"),
        benchmark_csv=str(INPUT_DIR / "HF index comparison.xlsx"),
    )


def extract_all_table_text(doc: Document) -> list:
    """Extract all table cell text as a list of (table_idx, row, col, text)."""
    results = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                results.append((t_idx, r_idx, c_idx, cell.text.strip()))
    return results


def extract_body_images(doc: Document) -> list:
    """Extract image info (body element index, rId) from inline images."""
    from docx.oxml.ns import qn
    results = []
    for idx, elem in enumerate(doc.element.body):
        inlines = elem.findall(".//" + qn("wp:inline"))
        for inline in inlines:
            blip = inline.find(".//" + qn("a:blip"))
            if blip is not None:
                rId = blip.get(qn("r:id"))
                if rId:
                    results.append((idx, rId))
    return results


def get_image_blob(doc: Document, rId: str) -> bytes:
    """Get the image bytes for a relationship ID."""
    if rId not in doc.part.rels:
        return b""
    rel = doc.part.rels[rId]
    return rel.target_part.blob


def _parse_percent(text: str):
    """Try to extract a numeric value from a percent string like '5.99%'."""
    m = re.match(r"^([+-]?\d+\.?\d*)%$", text.strip())
    return float(m.group(1)) if m else None


def compare_documents(generated_path: str, expected_path: str, numeric_atol: float = 0.05):
    """Compare two DOCX documents and return differences.

    Numeric percent cells that differ by less than *numeric_atol* (absolute)
    are treated as equal to tolerate minor benchmark rounding variance.
    """
    gen_doc = Document(generated_path)
    exp_doc = Document(expected_path)

    gen_cells = extract_all_table_text(gen_doc)
    exp_cells = extract_all_table_text(exp_doc)

    diffs = []

    # Compare table cell counts
    if len(gen_cells) != len(exp_cells):
        diffs.append(
            f"Cell count mismatch: generated={len(gen_cells)}, expected={len(exp_cells)}"
        )

    # Compare each cell
    for gen, exp in zip(gen_cells, exp_cells):
        t_idx, r_idx, c_idx = gen[:3]
        gen_text = gen[3]
        exp_text = exp[3]
        if gen_text != exp_text:
            # Allow small numeric differences in percent values
            gen_val = _parse_percent(gen_text)
            exp_val = _parse_percent(exp_text)
            if gen_val is not None and exp_val is not None and abs(gen_val - exp_val) < numeric_atol:
                continue
            diffs.append(
                f"Table {t_idx}, Row {r_idx}, Col {c_idx}: "
                f"generated='{gen_text}' expected='{exp_text}'"
            )

    # Compare images
    gen_images = extract_body_images(gen_doc)
    exp_images = extract_body_images(exp_doc)

    if len(gen_images) != len(exp_images):
        diffs.append(
            f"Image count mismatch: generated={len(gen_images)}, expected={len(exp_images)}"
        )

    # Compare image content by hash
    for i, (gen_img, exp_img) in enumerate(zip(gen_images, exp_images)):
        gen_blob = get_image_blob(gen_doc, gen_img[1])
        exp_blob = get_image_blob(exp_doc, exp_img[1])
        gen_hash = hashlib.md5(gen_blob).hexdigest()
        exp_hash = hashlib.md5(exp_blob).hexdigest()
        if gen_hash != exp_hash:
            diffs.append(
                f"Image {i} at body element {gen_img[0]} differs "
                f"(generated hash={gen_hash[:8]}, expected hash={exp_hash[:8]})"
            )

    return diffs


@pytest.mark.parametrize("lang", ["en", "cn"])
def test_generate_factsheet(funds, lang, tmp_path):
    """Generate factsheet and compare with expected outcome."""
    spec_path = str(TEMPLATE_DIR / f"docx_spec_{lang}.json")
    output_path = str(tmp_path / f"RDGFF_Factsheet_2026-01_{lang}.docx")
    expected_path = str(EXPECTED_DIR / f"RDGFF Factsheet_Jan_2026_{lang}.docx")

    # Verify expected file exists
    assert Path(expected_path).exists(), f"Expected outcome file not found: {expected_path}"

    # Generate
    result_path = generate_factsheet(
        funds=funds,
        spec_path=spec_path,
        report_month="2026-01",
        output_path=output_path,
    )

    assert Path(result_path).exists(), "Generated factsheet was not created"

    # Compare
    diffs = compare_documents(str(result_path), expected_path)

    if diffs:
        msg = f"\n{len(diffs)} difference(s) found in {lang} factsheet:\n"
        msg += "\n".join(f"  - {d}" for d in diffs)
        pytest.fail(msg)
